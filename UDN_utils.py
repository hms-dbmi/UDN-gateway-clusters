### All functions used to perform the UDN analysis 

"""TODO write description
    TODO comment all functions
"""

import argparse
import sys
import numpy as np
import matplotlib.pyplot as plt
from scipy.stats import ttest_ind
import scipy.stats as st
import scipy.sparse as sp
from scipy.stats import fisher_exact
import networkx as nx
from community import community_louvain
from scipy.stats import kruskal
import seaborn as sns
import collections as collec
import os
import xml.etree.ElementTree as ET
import operator
import pandas
import csv
from scipy.stats import mannwhitneyu
from sklearn.metrics.pairwise import pairwise_distances
from docx import Document
from docx.shared import Inches
import ast

import PicSureHpdsLib
import PicSureClient


## General methods

def get_CI(a):
    """Returns the 95% confidence interval for a list/array a
    Parameters: a (list): list or array we want the CI for
    Returns: (tuple) with 95% confidence interval
    """
    return st.t.interval(0.95, len(a)-1, loc=np.mean(a), scale=st.sem(a,nan_policy='omit'))

def get_data_df(column_head, resource):
    """Enables the user to download the data as a pandas dataframe indexed by UDN IDs (through API)
    Parameters : column_head (str): name of the header that will be selected. For example, if the columns that 
                                should be selected contain "this string", then column_head="this string".
                resource : picsure api resource
    Returns: df (pandas.core.Dataframe): dataframe indexed by UDN IDs of the selected columns
    """
    dictionary=resource.dictionary().find(column_head)

    query=resource.query()
    query.select().add(dictionary.keys())
    query.select().add('\\000_UDN ID\\')

    df=query.getResultsDataFrame()
    df.set_index("\\000_UDN ID\\", inplace=True)

    query.select().clear()

    return df

def create_table(n_ad,a_or_p,clusters_un,avg_HPO_clusters,CI_HPO_clusters,
                    gender_distrib,OR_diag,IC_OR,avg_onset,CI_onset,avg_UDN_eval,
                    CI_UDN_eval,docname):
    """Creates a word document with automatically rendered table of cluster characteristics
    Parameters: n_ad (int): number of clusters 
                a_or_p (str): "A" or "P", changes the display 
                clusters_un (dict): dictionnary with cluster as key and list of UDN IDs as value
                avg_HPO_clusters (dict): dictionnary with cluster as key and avg HPO per patient as value
                CI_HPO_clusters (dict): dictionnary with cluster as key and tuple of 95% CI for avg HPO as value
                gender_distrib (dict): dictionnary with cluster as key and count of female/male as value
                OR_diag (dict): dictionnary with cluster as key and OR as value
                IC_OR (dict): dictionnary with cluster as key and tuple of 95% CI for OR as value
                avg_onset (dict): dictionnary with cluster as key and avg onset as value
                CI_onset (dict): dictionnary with cluster as key and tuple of 95% CI for onset as value
                avg_UDN_eval (dict): dictionnary with cluster as key and avg UDN eval as value
                CI_UDN_eval (dict): dictionnary with cluster as key and tuple of 95% CI for avg UDN eval as value
                docname (str): name of document to save
    Returns: None
    Saves a word document with table of cluster characteristics
                
    """
    document = Document()

    document.add_heading('Tables'+docname, 0)

    table = document.add_table(rows=1, cols=n_ad+1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Clusters'
    for i in range(1,n_ad+1):
        hdr_cells[i].text = "Cluster C"+str(i)+a_or_p
    row_cells = table.add_row().cells
    row_cells[0].text = "# of patients per cluster"
    for i in range(1,n_ad+1):
        row_cells[i].text = str(len(clusters_un[i-1]))
    row_cells = table.add_row().cells
    row_cells[0].text = "Female:male ratio"
    for i in range(1,n_ad+1):
        row_cells[i].text = \
            str(int(np.round_(gender_distrib[i-1]["Female"]*10/gender_distrib[i-1]["Male"])))+":10"
    row_cells = table.add_row().cells
    row_cells[0].text = "Avg # of HPO terms per patient"
    for i in range(1,n_ad+1):
        row_cells[i].text = str(np.round_(avg_HPO_clusters[i-1],decimals=1))+ \
            " (95% CI: "+str(np.round_(CI_HPO_clusters[i-1][0],decimals=1))+ \
                " - "+str(np.round_(CI_HPO_clusters[i-1][1],decimals=1))+")"
    row_cells = table.add_row().cells
    row_cells[0].text = "Odds ratio diagnosed"
    for i in range(1,n_ad+1):
        row_cells[i].text = str(np.round_(OR_diag[i-1],decimals=1))+ \
            " (95% CI: "+str(np.round_(IC_OR[i-1]["low"],decimals=1))+ \
                " - "+str(np.round_(IC_OR[i-1]["up"],decimals=1))+")"
    row_cells = table.add_row().cells
    row_cells[0].text = "Average age at onset in y"
    for i in range(1,n_ad+1):
        row_cells[i].text = str(np.round_(avg_onset[i-1],decimals=1))+ \
            " (95% CI: "+str(np.round_(CI_onset[i-1][0],decimals=1))+ \
                " - "+str(np.round_(CI_onset[i-1][1],decimals=1))+")"
    row_cells = table.add_row().cells
    row_cells[0].text = "Average age at UDN evaluation in y"
    for i in range(1,n_ad+1):
        row_cells[i].text = str(np.round_(avg_UDN_eval[i-1],decimals=1))+ \
            " (95% CI: "+str(np.round_(CI_UDN_eval[i-1][0],decimals=1))+ \
                " - "+str(np.round_(CI_UDN_eval[i-1][1],decimals=1))+")"
    document.add_page_break()
    document.save(docname+'.docx')

def create_stat_table(kr_HPO_ad,kr_HPO_ped,kr_UDN_ad,kr_UDN_ped,kr_onset_ad,kr_onset_ped,docname):
    """Creates word document to save Kruskal Wallis results for clusters
    Parameters: kr_HPO_ad (tuple): Kruskal Wallis results of HPO for adult clusters
                kr_HPO_ped (tuple): Kruskal Wallis results of HPO for pediatric clusters
                kr_UDN_ad (tuple): Kruskal Wallis results of UDN eval for adult clusters
                kr_UDN_ped (tuple): Kruskal Wallis results of UDN eval for pediatric clusters
                kr_onset_ad (tuple): Kruskal Wallis results of onset age for adult clusters
                kr_onset_ped (tuple): Kruskal Wallis results of onset age for pediatric clusters
                docname (str): name of doc to save
    Returns: None
    Saves the word document with docname
    """
    document = Document()

    document.add_heading('Tables stats '+docname, 0)

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Variable'
    hdr_cells[1].text = 'Kruskal-Wallis H index and p-value'
    row_cells = table.add_row().cells
    row_cells[1].text = "Adult"
    row_cells[2].text = "Pediatric"
    row_cells = table.add_row().cells
    row_cells[0].text = "Avg # of HPO terms per patient"
    row_cells[1].text = "H = "+str(np.round_(kr_HPO_ad[0],decimals=1))+ \
        " , p = "+str(np.format_float_scientific(kr_HPO_ad[1],precision=2))
    row_cells[2].text = "H = "+str(np.round_(kr_HPO_ped[0],decimals=1))+ \
        " , p = "+str(np.format_float_scientific(kr_HPO_ped[1],precision=2))
    row_cells = table.add_row().cells
    row_cells[0].text = "Average age at onset in y"
    row_cells[1].text = "H = "+str(np.round_(kr_onset_ad[0],decimals=1))+ \
        " , p = "+str(np.format_float_scientific(kr_onset_ad[1],precision=2))
    row_cells[2].text = "H = "+str(np.round_(kr_onset_ped[0],decimals=1))+ \
        " , p = "+str(np.format_float_scientific(kr_onset_ped[1],precision=2))
    row_cells = table.add_row().cells
    row_cells[0].text = "Average age at UDN evaluation in y"
    row_cells[1].text = "H = "+str(np.round_(kr_UDN_ad[0],decimals=1))+ \
        " , p = "+str(np.format_float_scientific(kr_UDN_ad[1],precision=2))
    row_cells[2].text = "H = "+str(np.round_(kr_UDN_ped[0],decimals=1))+ \
        " , p = "+str(np.format_float_scientific(kr_UDN_ped[1],precision=2))
    document.add_page_break()

    document.save(docname+'.docx')

## Plotting functions

def show_distrib_HPO(HPO_list,name):
    """Plots the distribution of count of HPO terms per patient
    Parameters : HPO_list (list): list of counts for each patient of HPO terms
                 name (str): string, title of the figure
    Returns : None
    Shows matplotlib plot of distribution of HPO
    """
    distrib=collec.Counter(HPO_list)
    X=[key for key in distrib.keys()]
    Y=[distrib[key] for key in distrib.keys()]
    plt.figure(figsize=(20,15))
    plt.plot(X,Y,"o")
    plt.xlabel("Number of HPO terms",fontsize=40)
    plt.ylabel("Count of patients",fontsize=40)
    plt.title(name,fontsize=50)
    plt.xticks(fontsize=40)
    plt.yticks(fontsize=40)
    plt.yscale("log")
    plt.xscale("log")
    plt.axes().set_ylim(None,200)
    plt.show()
    plt.savefig("HPO_terms_log")

def show_age_distrib(demographics):
    """Show the age distribution in the network
    Parameters: demographics (pandas.core.DataFrame): olumns containing age at symptom onset
    Returns: None
    Shows the age distribution as a plot
    """
    X=list(collec.Counter(demographics["\\00_Demographics\\Age at symptom onset in years\\"].fillna(0)))
    Y=[collec.Counter(demographics["\\00_Demographics\\Age at symptom onset in years\\"])[i] for i in X]
    plt.figure(figsize=(20,20))
    plt.plot(X,Y)
    plt.title("Age at symptom onset (in y) distribution in UDN")
    plt.xlabel("Age at symptom onset (in y)")
    plt.ylabel("Count of patients")
    plt.show()

def plot_distribution_genomic_data(genomic_data,namefile,var_or_gene):
    """Show the distribution of counts of candidate genes or variant per patient in the UDN database
    Parameters: genomic_data (dict): dictionary, with UDN ID as key and list with dictionaries as value, dict contaning characteristics
                              of the considered genomic data
                namefile (str): file of the name to save the figure in 
                var_or_gene (str): "variants" if variants is considered, "genes" else
    Returns: None
    Show the distribution in a scatter plot and the counter, as well as total number of candidate genes/variants
    """
    count_gene_per_patient=collec.Counter([len(genomic_data[patient]) for patient in genomic_data])
    X_gene=list(count_gene_per_patient)
    Y_gene=[count_gene_per_patient[ct] for ct in X_gene]
    print("Number of total candidate ",var_or_gene," : ",np.sum([X_gene[i]*Y_gene[i] for i in range(len(X_gene))]))
    plt.figure(figsize=(10,5))
    plt.plot(X_gene,Y_gene,"o")
    plt.xticks(np.arange(0,18))
    plt.title("Distribution of number of candidate "+var_or_gene+" per patient")
    plt.xlabel("Number of candidate "+var_or_gene)
    plt.ylabel("Count of patients")
    plt.savefig(namefile,bbox_inches="tight",dpi=300)
    plt.show()

def heatmap_phen(clusters_un,phen_ranked,ind_groups,ad_or_ped,nb_phen,figsize,vmin,vmax,figname):
    """Displays heatmap of phenotype enrichment analysis for each cluster with analyzed composition
    Parameters: clusters_un (dict): dictionary with cluster number as key and list of patients in the 
                                    cluster as value
                phen_ranked (dict): dictionary with cluster number as key, two arrays as value, 
                                    one with list of phenotypes ranked according to composition, 
                                    second with composition of each phenotype
                ind_groups (list): indices to take into consideration
                ad_or_ped (str): "adult" or "pediatric", changes the display
                nb_phen (int): number of best phen to display
                figsize (int): size of the figure displayed
                vmin (int): minimum value for the heatmap (here percentage)
                vmax (int): max value for the heatmap (here percentage)
                figname (str): name under which you save the heatmap

    Returns: None
    Shows the heatmap of phenotype enrichment analysis for each cluster
    """
    if ad_or_ped=="adult":
        cluster_list=["Cluster C"+str(cluster+1)+"A, N="+ \
            str(len(clusters_un[cluster])) for cluster in ind_groups]
    elif ad_or_ped=="pediatric":
        cluster_list=["Cluster C"+str(cluster+1)+"P, N="+ \
            str(len(clusters_un[cluster])) for cluster in ind_groups]
    list_phen_max=[]
    for cluster in ind_groups:
        i,j=0,0
        while j<nb_phen:
            if not(phen_ranked[cluster][0][i]) in list_phen_max:
                list_phen_max.append(phen_ranked[cluster][0][i])
                j+=1
            i+=1
    heatmap_mat=[[] for i in range(len(list_phen_max))]
    for i,phen in enumerate(list_phen_max):
        for cluster in ind_groups:
            if phen in phen_ranked[cluster][0]:
                indphen=np.where(phen_ranked[cluster][0]==phen)[0][0]
                heatmap_mat[i].append(phen_ranked[cluster][1][indphen]*100)
            else:
                heatmap_mat[i].append(0)
    sns.set()
    fig,ax=plt.subplots(figsize=(figsize,figsize))
    sns.heatmap(heatmap_mat,cbar=True,cmap="YlGnBu",xticklabels=cluster_list,
                    yticklabels=list_phen_max,ax=ax,vmin=vmin,vmax=vmax)
    plt.ylabel("Phenotypes")
    plt.savefig(figname+".svg",bbox_inches="tight",dpi=350)
    plt.show()

def heatmap_real(count_dg_real,clusters_un,nb_outliers,figsize,vmin,vmax,figname):
    """Show the heatmap of cluster according to disease group composition
    Parameters: count_dg_real (dict): dictionary with cluster number as key, dictionary as value, with disease group as key and count of disease
                               group
                clusters_un (dict): dictionary with cluster number as key, list of patients as value
                nb_outliers (int): nb of outliers in network (adult or pediatric)
                figsize (int): size of the figure
                vmin (int): minimum value for the heatmap (here percentage)
                vmax (int): max value for the heatmap (here percentage)
                figname (str): name of figure to save
    Returns: None
    Shows the heatmap associated to the composition in disease group for each cluster
    """
    type_disease={'chromosomal_anomalies': 0, 'rare_abdominal_surgical_diseases': 1, 'rare_allergic_diseases': 2, 
                  'rare_bone_diseases': 3, 'rare_cardiac_diseases': 4, 'rare_cardiac_malformations': 5, 
                 'rare_circulatory_system_diseases': 6, 'rare_developmental_anomalies_during_embryogenesis': 7, 
                  'rare_endocrine_diseases':  8, 'rare_eye_diseases': 9, 'rare_gastroenterological_diseases': 10,
                 'rare_genetic_diseases': 11, 'rare_gynaecological_and_obstetric_diseases': 12, 'rare_haematological_diseases': 13,
                 'rare_hepatic_diseases': 14, 'rare_immunological_diseases': 15, 'rare_inborn_errors_of_metabolism': 16,
                 'rare_infectious_diseases': 17, 'rare_infertility': 18, 'rare_intoxications': 19, 'rare_neoplastic_diseases': 20,
                 'rare_neurological_diseases': 21, 'rare_odontological_diseases': 22, 'rare_otorhinolaryngological_diseases': 23,
                 'rare_renal_diseases': 24, 'rare_respiratory_diseases': 25, 'rare_rheumatological_diseases_of_childhood': 26,
                 'rare_skin_diseases': 27, 'rare_sucking_swallowing_disorders': 28, 'rare_surgical_maxillo_facial_diseases': 29,
                 'rare_surgical_thoracic_diseases': 30, 'rare_systemic_and_rheumatological_diseases': 31, 
                  'rare_teratologic_disorders': 32, 'rare_urogenital_diseases': 33}
    heatmap_mat=[[0 for i in range(len(type_disease))] for j in range(len(count_dg_real))]
    cluster_list=np.sort([cluster for cluster in count_dg_real])
    yl=["" for i in range(len(type_disease))]
    for dis in type_disease:
        listname=dis.split(".")[0].split("_")
        name=""
        for string in listname:
            name+=string+" "
        yl[type_disease[dis]]=name
    for cluster in count_dg_real:
        for dg in count_dg_real[cluster]:
            index=cluster_list.tolist().index(cluster)
            heatmap_mat[index][type_disease[dg]]=count_dg_real[cluster][dg]
    cluster_list=["Cluster C"+str(cluster+1)+", N="+str(len(clusters_un[cluster])) for cluster in count_dg_real]
    cluster_list[len(cluster_list)-1]="Outliers, N={}".format(nb_outliers)
    sns.set()
    fig,ax=plt.subplots(figsize=(figsize*1.3,figsize//1.5))
    ax=sns.heatmap(heatmap_mat,cbar=True,xticklabels=yl,yticklabels=cluster_list,vmin=vmin,vmax=vmax)
    ax.set_xticks(np.arange(len(type_disease))+0.5)
    ax.set_yticks(np.arange(len(count_dg_real))+0.5)
    plt.setp(ax.get_xticklabels(), rotation=45, ha="right",
         rotation_mode="anchor", va="top")
    plt.setp(ax.get_yticklabels(), rotation=0, ha="right",
            rotation_mode="default", va="center_baseline")
    plt.xlabel("Disease subgroup")
    plt.savefig(figname+".png",bbox_inches="tight",dpi=500)
    plt.show()
